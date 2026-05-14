"""
Build a Word document containing the full speech Cagri will deliver to students,
slide by slide. Designed to be read straight off a tablet by a presenter who has
not memorized the slides in advance.

Each slide section has three labeled blocks:
  ON SCREEN  — what the slide currently shows (so the presenter knows the context)
  SAY        — the spoken script, written in plain conversational English
  DO         — physical actions to take (point at slide, wait, click, walk around)
"""
import os
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

HERE = os.path.dirname(os.path.abspath(__file__))
ROOT = os.path.dirname(HERE)
OUT_PATH = os.path.join(ROOT, 'WSU-Workshop-Speech.docx')

doc = Document()

for section in doc.sections:
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)
    section.left_margin = Inches(0.9)
    section.right_margin = Inches(0.9)

style = doc.styles['Normal']
style.font.name = 'Georgia'
style.font.size = Pt(13)


# ============ helpers ============

def add_title_page():
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('From Black Boxes to Glass Boxes')
    r.font.size = Pt(28); r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Read-Off-The-Page Presentation Script')
    r.font.size = Pt(18); r.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    doc.add_paragraph()
    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Washington State University · Data and Analytics Breakout')
    r.font.size = Pt(14); r.italic = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('May 15, 2026 · 09:00 to 11:45 (165 minutes)')
    r.font.size = Pt(14); r.italic = True

    doc.add_paragraph()

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('Cagri Temel')
    r.font.size = Pt(16); r.font.bold = True

    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = p.add_run('CTO, Hezarfen LLC  ·  IEEE Senior Member')
    r.font.size = Pt(12); r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph(); doc.add_paragraph(); doc.add_paragraph()

    p = doc.add_paragraph()
    r = p.add_run('How to use this script')
    r.font.size = Pt(14); r.font.bold = True

    intro = (
        "This script is designed so you can read it straight off a tablet without "
        "needing to study the slides in advance. Each slide section has three blocks:\n\n"
        "ON SCREEN  — what the slide on the projector is showing right now.\n"
        "SAY        — the words you speak. Read them out loud at a normal pace.\n"
        "DO         — physical actions: point at the slide, wait, click, walk around.\n\n"
        "If you just read every SAY line in order, you have a complete workshop. "
        "The ON SCREEN lines remind you what students are looking at. "
        "The DO lines remind you when to pause, point, or move."
    )
    p = doc.add_paragraph(intro)
    p.paragraph_format.space_after = Pt(6)

    doc.add_page_break()


def add_section_heading(text, time_marker=None):
    p = doc.add_paragraph()
    r = p.add_run(text.upper())
    r.font.size = Pt(20); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    p.paragraph_format.space_before = Pt(12)
    p.paragraph_format.space_after = Pt(4)
    if time_marker:
        p = doc.add_paragraph()
        r = p.add_run(time_marker)
        r.font.size = Pt(11); r.font.italic = True
        r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)
        p.paragraph_format.space_after = Pt(12)
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement('w:pBdr')
    bottom = OxmlElement('w:bottom')
    bottom.set(qn('w:val'), 'single')
    bottom.set(qn('w:sz'), '6')
    bottom.set(qn('w:space'), '1')
    bottom.set(qn('w:color'), 'cccccc')
    pBdr.append(bottom)
    pPr.append(pBdr)


def add_slide_header(slide_num, slide_title, time_clock):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(20)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(f'Slide {slide_num}  ·  {slide_title}')
    r.font.size = Pt(15); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r = p.add_run(f'        {time_clock}')
    r.font.size = Pt(11); r.font.italic = True
    r.font.color.rgb = RGBColor(0x88, 0x88, 0x88)


def add_label(label, body, color):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.4
    r = p.add_run(f'{label}  ')
    r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = color
    r = p.add_run(body)
    r.font.size = Pt(13)


def screen(text):
    add_label('ON SCREEN', text, RGBColor(0x66, 0x66, 0x66))


def say(text):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.line_spacing = 1.45
    p.paragraph_format.left_indent = Inches(0.3)
    r = p.add_run('SAY     ')
    r.font.size = Pt(11); r.font.bold = True
    r.font.color.rgb = RGBColor(0x1F, 0x4E, 0x79)
    r = p.add_run(text)
    r.font.size = Pt(13)


def do(text):
    add_label('DO         ', text, RGBColor(0xB7, 0x67, 0x00))


# ============ content ============

add_title_page()

add_section_heading('Block 1 — Lecture (slides 1 to 19)',
                    'Running time 09:00 to 09:30 · 30 minutes')


# ---- 1
add_slide_header(1, 'Cover', '09:00')
screen("Workshop title 'From Black Boxes to Glass Boxes' with your name and IEEE badge.")
say("Good morning everyone. I'm Cagri Temel, CTO at Hezarfen LLC. "
    "For the next two hours and forty-five minutes you're going to train a real machine learning "
    "model on real NASA data, break it on purpose, and then explain how it makes its decisions.")
say("By the end of today you will have a working model on your laptop, you will know which sensor "
    "failed on an aircraft engine, and many of you will have your first open-source contribution on GitHub. "
    "Let's get started.")

# ---- 2
add_slide_header(2, 'Who I am', '09:01')
screen("Short bio: CTO Hezarfen, IEEE Senior Member, author of the neural-trees Python package.")
say("Quick about me. I'm CTO of Hezarfen, a company that builds explainable AI for regulated industries — "
    "banks, insurers, healthcare, aviation. I'm an IEEE Senior Member. I also maintain a Python package "
    "called neural-trees, and that's the tool we'll use today.")
say("You don't need to memorize any of this. The one line I want you to remember from this slide is at "
    "the bottom: in two hours and forty-five minutes, you build, you break, and you explain a real model.")

# ---- 3
add_slide_header(3, "Today's Mission", '09:02')
screen("Big title 'Data-driven decisions, when the cost of being wrong is high' over a dark slide.")
say("Today we're looking at one question. How do you make data-driven systems trustworthy in places "
    "where being wrong is expensive — where it costs lives, or money, or certifications?")
say("Our case is a turbofan engine. The kind that powers passenger jets. It needs to predict its own "
    "failure before it fails. By the end of the session you'll see why this is harder than it looks, "
    "and why being accurate is not enough.")

# ---- 4
add_slide_header(4, 'Hook', '09:04')
screen("Just two short sentences on a dark slide: 'A turbofan engine. When will it fail?'")
say("A turbofan engine. When will it fail?")
do("Pause two seconds. Don't say anything. Let the question land.")
say("If you knew with certainty when this engine would fail, what would you do differently than what "
    "airlines do today?")
do("Don't call on anyone. Move to the next slide after a beat.")

# ---- 5
add_slide_header(5, 'Stakes', '09:05')
screen("Three bullet lines about predicting too late, too early, or accurately-but-unexplainably.")
say("There are three ways to be wrong, and all three are expensive.")
do("Point at the first bullet.")
say("Predict too late. The engine fails mid-flight. People die.")
do("Point at the second bullet.")
say("Predict too early. A perfectly good engine gets pulled. A hundred million dollars wasted.")
do("Point at the third bullet.")
say("Predict accurately but you can't explain how. The FAA refuses to certify your system. "
    "Best model in the world, can't ship it, can't fly it.")
say("The whole point: accuracy alone is not enough. You also need to defend the prediction.")

# ---- 6
add_slide_header(6, 'NASA CMAPSS', '09:08')
screen("Bullet list describing the NASA CMAPSS dataset: 100 engines, 21 sensors, public.")
say("The dataset we'll use is called NASA CMAPSS. It's open. It's free. It's been the standard "
    "predictive-maintenance benchmark for about fifteen years. A hundred turbofan engines, each one "
    "simulated all the way to failure. Twenty-one sensors per timestep. Each engine logs hundreds "
    "of flight cycles before it dies.")
say("You will load this data on your own laptop in about twenty minutes.")

# ---- 7
add_slide_header(7, 'Remaining Useful Life', '09:10')
screen("Definition of RUL with an ASCII strip showing RUL counting down to 0 (failure).")
say("The thing we're going to predict is called RUL — Remaining Useful Life. The number of flight "
    "cycles left before the engine fails.")
say("So if an engine fails at cycle 192: at cycle 100 its RUL is 92. At cycle 190 its RUL is 2. "
    "At failure, RUL is zero.")
say("We cap RUL at 125 because predicting further out than that is not useful. Past a hundred and "
    "twenty-five cycles, you're just guessing about the future.")

# ---- 8
add_slide_header(8, '21 Sensors', '09:12')
screen("A table of the 21 sensor channels with three rows highlighted in blue: sensors 7, 11, 14.")
say("For every cycle we get twenty-one sensor readings. You don't need to memorize them.")
say("But three of them I want you to notice now, because they will come back.")
do("Point at sensor 7 in the table.")
say("Sensor seven — high-pressure compressor outlet pressure.")
do("Point at sensor 11.")
say("Sensor eleven — static pressure.")
do("Point at sensor 14.")
say("Sensor fourteen — corrected core speed.")
say("These three are the model's favorites. They will show up again in Activity 2 — when one of them "
    "gets attacked.")

# ---- 9
add_slide_header(9, 'LSTM Result', '09:14')
screen("A small Python snippet defining an LSTM, with RMSE 15.49 and R-squared 0.855 printed below.")
say("Here's the modern approach. We throw an LSTM at it. Two-layer recurrent network. After training, "
    "RMSE fifteen-point-four-nine cycles. R-squared 0.855.")
say("That's pretty good. About fifteen cycles off, on average.")
say("So. We ship it to the FAA, right?")

# ---- 10
add_slide_header(10, 'Regulator Problem', '09:16')
screen("A huge red 'No.' on a dark slide, then a quote about why regulators reject black-box models.")
say("No.")
do("Pause two seconds. Don't fill the silence.")
say("Regulators don't ask 'how accurate is your model?' They ask 'why did your model say that?' "
    "If you can't answer that in a way a non-engineer can follow, your model doesn't ship.")
say("FAA for aviation. FDA for medical devices. OCC for banking. EU AI Act for anything that touches "
    "a European customer. This is not theory. This is what actually blocks AI from being deployed in "
    "regulated industries.")

# ---- 11
add_slide_header(11, 'Black Box vs Glass Box', '09:18')
screen("Two cards side by side: red 'Black Box' (LSTM, accurate, unexplainable) and green 'Glass Box' "
       "(decision tree, explainable, less accurate).")
say("You're stuck between two worlds.")
do("Point at the red card.")
say("On one side: black-box models. LSTMs, transformers. High accuracy. You can't explain why.")
do("Point at the green card.")
say("On the other side: glass-box models. Decision trees, linear regression. Fully explainable. "
    "Lower accuracy.")
say("Regulators love glass boxes. Engineers prefer black boxes. The question today is — can we have both?")

# ---- 12
add_slide_header(12, 'Classic Decision Tree', '09:20')
screen("An ASCII tree diagram with three splits and three leaf labels (Healthy, Caution, Critical).")
say("You've all seen a decision tree before. Beautiful structure. Sensor eleven is above the threshold, "
    "go left. Below, go right. It's the most explainable model that exists.")
say("But classic decision trees often have worse accuracy than neural networks. Why? Because the splits "
    "are hard. A point-zero-one difference in a sensor value flips the entire path. That brittleness is "
    "what limits accuracy.")

# ---- 13
add_slide_header(13, 'Hard Split', '09:22')
screen("Five lines of Python code defining a hard split using if-else, returning the strings 'right' or 'left'.")
say("This is the code for a hard split. You are NOT typing this. It is here so you see the contrast.")
say("If x is bigger than the threshold, return right. Otherwise, return left. The output is a string. "
    "Discrete. No gradient.")
say("No gradient means no backpropagation. You build the tree node by node, greedily. You cannot train "
    "it end to end like a neural network.")

# ---- 14
add_slide_header(14, 'Soft Split', '09:23')
screen("Same five-line function but with the if-else replaced by a sigmoid, returning a probability in [0,1].")
say("Same function. One substitution. The if-else becomes a sigmoid.")
say("Now the output is a probability. Continuous. And because it's continuous, gradient flows through it.")
say("That single change is what makes everything we do today possible. A tree you can train with "
    "backpropagation, just like a neural network.")

# ---- 15
add_slide_header(15, 'Soft Decision Tree', '09:25')
screen("An ASCII visualization of a full soft tree with sigmoid splits at each node and leaf distributions.")
say("Here's what the whole tree looks like. Every internal node is a sigmoid split. Every leaf has "
    "a class distribution — Critical, Caution, Healthy probabilities.")
say("The final prediction is a weighted sum of all the leaves, weighted by the probability of reaching "
    "each leaf. It is mathematically smooth and fully differentiable end to end.")

# ---- 16
add_slide_header(16, 'Best of Both', '09:26')
screen("Two cards: green 'From the tree' (path traceability) and blue 'From the network' (end-to-end training).")
say("Here's the win.")
do("Point at the green 'From the tree' card.")
say("From the tree, you keep path traceability, per-feature thresholds, human-readable rules.")
do("Point at the blue 'From the network' card.")
say("From the network, you get end-to-end training and the accuracy that comes with it.")
say("You get both. That's why soft decision trees matter for regulated industries.")

# ---- 17
add_slide_header(17, 'Three Properties', '09:27')
screen("Three numbered properties: explainability, noise robustness, sensor failure tolerance.")
say("Three properties make this approach work for safety-critical AI.")
do("Point at property 1.")
say("Explainability. Every prediction comes with the path that produced it.")
do("Point at property 2.")
say("Noise robustness. Soft splits don't break when a sensor wobbles a little.")
do("Point at property 3.")
say("Sensor failure tolerance. When you train with channel-level dropout, the model learns to operate "
    "even when some sensors are missing.")
say("You'll see all three in the next ninety minutes. Not on a slide. On your own laptop.")

# ---- 18
add_slide_header(18, 'Paper Numbers', '09:28')
screen("A four-row table comparing LSTM, Random Forest, and the Temporal Neural Tree on clean RMSE and "
       "degradation under 30% missing sensors.")
say("These are the numbers from my paper.")
do("Point at the second column.")
say("On clean CMAPSS data, an LSTM gets RMSE fifteen-point-four-nine cycles. The Temporal Neural Tree "
    "gets fifteen-point-seven-eight. Essentially equal.")
do("Point at the third column.")
say("When thirty percent of the sensors are missing — which happens in real aircraft, sensors break — "
    "the LSTM degrades by eighty-nine percent. The Temporal Neural Tree degrades by only seventeen percent.")
say("Five times more robust. And it's fully explainable. That's the headline.")

# ---- 19
add_slide_header(19, 'neural-trees Package', '09:29')
screen("Big title 'neural-trees' with a 'pip install neural-trees' command and four bullet points.")
say("The tool you'll use is called neural-trees. I wrote it. On PyPI. MIT licensed. Open source.")
say("The API follows scikit-learn — fit, predict, predict_proba, score. If you've used scikit-learn, "
    "you already know how to use this. The backend is PyTorch. You don't need a GPU; CPU is fine for "
    "everything we do today.")
say("At the end of the workshop, you'll get a chance to contribute back to it.")


doc.add_page_break()

# ============ BLOCK 2 ============
add_section_heading('Block 2 — Activity 1: Train Your Own Neural Tree',
                    'Running time 09:30 to 10:25 · 55 minutes')

# ---- 20
add_slide_header(20, 'Activity 1 Intro', '09:30')
screen("Large title 'Train Your Own Neural Tree' with a 55-minute label.")
say("Okay, we're moving from lecture to hands-on. The next fifty-five minutes are yours.")
say("You're going to train your own neural tree on the NASA data we just talked about. I'm going to "
    "walk around. If you get stuck, raise a hand and I'll come help.")
say("There are eleven steps. Each step is on its own slide. You paste the code, run it, check that the "
    "output matches mine. Ready? Let's go.")

# ---- 21
add_slide_header(21, 'QR Code', '09:31')
screen("A large QR code on a white slide, and the workshop GitHub URL below it.")
say("Pull out your phones and scan this QR code. It opens our workshop landing page.")
say("Tap Activity 1. That opens a Google Colab notebook on your laptop. Sign in with your Gmail. "
    "Then click 'Copy to Drive' at the top so you have your own editable copy.")
do("Watch the room for two minutes. Do not advance until you can see most laptops on Colab. "
   "Help anyone who looks lost.")

# ---- 22
add_slide_header(22, 'How a Colab Notebook Works', '09:33')
screen("Two columns: left explains cells, right explains run-status icons. A red warning at the bottom "
       "about running cells in order.")
say("Thirty seconds on how Colab works, for anyone who's never used a notebook.")
do("Point at the left column.")
say("Your notebook is a list of cells. Each gray box is one cell. Click into a cell, paste code with "
    "Command-V or Control-V, then press Shift+Enter to run. Output shows up right below.")
do("Point at the right column.")
say("While a cell is running, you'll see a star in brackets next to it. When it finishes, the star turns "
    "into a number. If you see a red box, that's an error. Always read the last line of red text — "
    "that's the real error message.")
say("Most important rule: run cells in order from top to bottom. Don't skip. If you do, you'll get "
    "'name is not defined' errors.")
say("Okay. Let's look at the next fifty-five minutes.")

# ---- 23
add_slide_header(23, 'Activity 1 Overview', '09:34')
screen("Numbered list of the 11 steps with rough time estimates. Step 5 is highlighted in blue.")
say("Quick overview. Six stages. We load CMAPSS, look at it, bin the labels into three classes, train "
    "the soft decision tree — that takes about a minute — evaluate it with a confusion matrix.")
do("Point at step 5.")
say("And then — the step that matters most — we traverse one prediction. We literally read out the rule "
    "the model used. If you remember one thing from this entire workshop, it should be that step. "
    "Let's start.")

# ---- 24
add_slide_header(24, 'Step 1 — Imports', '09:35')
screen("STEP 1 / 11 tag, the title 'Imports'. Left side: a code block with eight import lines. "
       "Right side: a green box showing 'All set.'")
say("First a quick reminder before Step 1. At the very top of your notebook there is a Setup cell. "
    "Run that one FIRST. It installs the neural-trees package and downloads the data files. "
    "Takes about twenty seconds. You should see 'Setup done.' before you do anything else.")
do("Wait until you can see most students have run the Setup cell.")
say("Okay, now Step 1. Imports.")
do("Point at the code block on the left.")
say("Copy this code from the slide. Paste it into the empty cell under the Step 1 heading. "
    "Press Shift+Enter.")
do("Point at the green output box on the right.")
say("You should see 'All set.' If you see an IndentationError, that means one of your lines has a stray "
    "space at the beginning. Click into the cell, hit Command-A, then Shift+Tab to clear it, then run "
    "again.")
do("Walk around. Three minutes. Move on when most of the room has 'All set.' on their screen.")

# ---- 25
add_slide_header(25, 'Step 2 — Load Data', '09:38')
screen("STEP 2 / 11. Code on left loads the CSV with pandas. Right side shows expected (20631, 26) -> 100 engines.")
say("Step 2. We load the CMAPSS data.")
do("Point at the code on the left.")
say("Paste this code into the next empty cell. Shift+Enter.")
do("Point at the green output box on the right.")
say("You should see the shape — twenty thousand six hundred and thirty-one rows by twenty-six columns — "
    "and 'one hundred engines'. Below that, the first five rows of engine one.")
say("Those numbers — temperatures, pressures, fan speeds — those are exactly what a flight data recorder "
    "logs in real life.")
do("Walk around. Two minutes.")

# ---- 26
add_slide_header(26, 'Step 3 — Compute RUL', '09:41')
screen("STEP 3 / 11. Code computes RUL. Right side shows a table with engine 1's first five rows all "
       "with RUL = 125.")
say("Step 3. We compute the label we're trying to predict.")
do("Point at the code.")
say("For each row, RUL equals the engine's maximum cycle minus the current cycle. We cap it at a "
    "hundred and twenty-five. Paste, run.")
do("Point at the output table.")
say("All five rows show RUL equal to a hundred and twenty-five — because the engine is brand new and "
    "the real RUL is much higher, but we capped it. That's expected.")
do("Two minutes. Move on.")

# ---- 27
add_slide_header(27, 'Step 4 — Plot Engine 1', '09:44')
screen("STEP 4 / 11. Code on left. Right side shows a 2x2 plot of sensors 2, 7, 11, 14 over engine 1's "
       "192-cycle life.")
say("Step 4. We visualize one engine.")
do("Point at the code.")
say("Paste, run. You'll get a 2-by-2 plot — four sensors for engine one across its entire life.")
do("Point at the plot on the right.")
say("Look carefully. Sensor two — temperature — goes up. Sensor eleven — pressure — goes up. Sensor "
    "fourteen — corrected core speed — goes up. But sensor seven — HPC outlet pressure — goes down.")
say("That's the engine wearing out. Temperatures rising, compressor pressure dropping. The data is "
    "telling us the story of failure. Now we have to teach a model to see it.")
do("Three minutes. Move on.")

# ---- 28
add_slide_header(28, 'Step 5 — Bin RUL', '09:49')
screen("STEP 5 / 11. Code defines bin_rul function. Right side shows class counts: 3000 Critical, "
       "5000 Caution, 12631 Healthy.")
say("Step 5. We turn RUL into a classification problem.")
say("In real aviation no one says 'this engine has forty-seven cycles left.' They say healthy, "
    "caution, or critical. So we bin RUL into three classes.")
do("Point at the code.")
say("Below thirty is Critical. Between thirty and eighty is Caution. Above eighty is Healthy. Paste, run.")
do("Point at the output.")
say("Roughly fourteen percent Critical, twenty-four percent Caution, sixty-one percent Healthy. The "
    "classes are imbalanced. That's reality. Most of the time, an engine is fine.")

# ---- 29
add_slide_header(29, 'Step 6 — Feature Matrix', '09:53')
screen("STEP 6 / 11. Code drops constant sensors and standardizes. Right side: 'X shape: (20631, 15) y shape: (20631,)'.")
say("Step 6. Build the feature matrix.")
do("Point at the code.")
say("We drop the six sensors that are constant — they carry no information. That leaves fifteen "
    "informative sensors. Then we standardize each column to mean zero and unit variance. Paste, run.")
do("Point at the output.")
say("You should see X shape twenty thousand by fifteen, y shape twenty thousand. Soft decision trees "
    "don't strictly require standardization, but the sigmoid splits behave much better when the input "
    "is on a consistent scale.")

# ---- 30
add_slide_header(30, 'Step 7 — Train / Test Split', '09:57')
screen("STEP 7 / 11. Code does an 80/20 stratified split. Right side: (16504, 15) and (4127, 15).")
say("Step 7. Standard train-test split. Eighty percent training, twenty percent test.")
do("Point at the code.")
say("Stratify equals y — that preserves the class ratio in both sets. Random state equals forty-two, "
    "so everyone in this room gets the exact same numbers. If yours are different, you typed something "
    "wrong. Paste, run.")
do("Point at the output.")
say("Sixteen thousand five hundred and four rows for training. Four thousand one hundred and twenty-seven "
    "for test.")

# ---- 31
add_slide_header(31, 'Step 8 — Train SDT', '10:00')
screen("STEP 8 / 11. Code instantiates SoftDecisionTree, calls fit, prints accuracy. Right side shows "
       "epoch-by-epoch loss going down and final 'Test accuracy: 0.841'.")
say("Step 8. We train the Soft Decision Tree. This is the slow one. About thirty to sixty seconds on CPU.")
do("Point at the code.")
say("Paste, run. While we wait, let me tell you what's happening: PyTorch is doing backpropagation on "
    "sixteen thousand rows for thirty epochs. Every epoch, the loss should come down a little.")
say("When it finishes you should see test accuracy around zero-point-eight-four. Anywhere between "
    "zero-point-eight-two and eight-six is normal — randomness varies a little across machines.")
do("Wait for the slowest student to finish. While waiting, ask: 'What do you think an LSTM would "
   "get on the same problem?' Most will guess around the same. Then say: 'Right. About the same. "
   "The point isn't accuracy. The point is what we can do with the model next.'")

# ---- 32
add_slide_header(32, 'Step 9 — Confusion Matrix', '10:05')
screen("STEP 9 / 11. Code on left. Right side: a blue heatmap showing the 3x3 confusion matrix and a "
       "classification report below.")
say("Step 9. We look at where the model is right and where it's wrong.")
do("Point at the code.")
say("Paste, run. You'll get a confusion matrix — three by three, color-coded — plus a per-class report.")
do("Point at the bottom-left cell of the heatmap.")
say("That cell should say zero. Out of every Healthy engine in the test set, the model wrongly flagged "
    "zero of them as Critical. We never raised a false alarm.")
do("Point at the top row.")
say("Out of six hundred Critical engines, we caught about four hundred and ninety. We missed some — but "
    "only five of those misses went all the way to Healthy.")
say("In aviation language: we don't miss many dangerous engines, and we don't waste healthy ones. "
    "That's a useful model.")

# ---- 33
add_slide_header(33, 'Step 10 — Split Weights', '10:10')
screen("STEP 10 / 11. Code prints dominant sensor per internal node. Right side: a list of 15 nodes "
       "with sensor names; s14 and s6 appear multiple times.")
say("Step 10. We peek inside the model.")
do("Point at the code.")
say("Paste, run. You'll see fifteen internal nodes and the sensor each one leans on most.")
do("Point at the output list on the right.")
say("Notice some sensors show up multiple times. Sensor fourteen and sensor six in particular. Those "
    "are the pillars of the model. If those sensors go bad, the model is in trouble.")
say("Remember sensor fourteen. It comes back in Activity 2.")

# ---- 34
add_slide_header(34, 'Step 11 — Traverse Prediction', '10:15')
screen("STEP 11 / 11. Code reads one prediction. Right side: four lines showing True=Caution, "
       "Predicted=Healthy, probabilities, root sensor=s7.")
do("Slow down. This is the most important step. Hit the pause button on tempo.")
say("Step 11. This is the step that matters most.")
do("Point at the code.")
say("We pick one specific test sample — sample seventeen — and we ask the model: what did you predict, "
    "with what confidence, and which sensor did you look at first? Paste, run.")
do("Point at each of the four output lines, one at a time.")
say("True class: Caution. Predicted class: Healthy. Probabilities: zero percent Critical, fourteen percent "
    "Caution, eighty-five percent Healthy. Root sensor: sensor seven — HPC outlet pressure.")
do("Look at the room. Now deliver the punchline.")
say("Now imagine you have to explain this prediction to a regulator. With this output, you can say: "
    "'My model classified this engine as Healthy because the static pressure at the high-pressure "
    "compressor outlet was above its decision threshold.' An LSTM cannot give you this answer. "
    "That is why we did all of this.")

# ---- 35
add_slide_header(35, 'Checkpoint', '10:23')
screen("Green slide. Title 'Checkpoint'. A short checklist of what students should have at this point.")
say("Okay. Take a breath. By now you should have three things on your screen.")
do("Point at the checklist.")
say("A trained Soft Decision Tree. Test accuracy above eighty percent. A printed decision path for "
    "engine seventeen.")
say("If you're missing any of those, find me during the break and we'll catch up. Ten-minute break. "
    "Get coffee, stretch. When you come back at ten thirty-five, we run the team competition.")


doc.add_page_break()


# ============ BLOCK 3 ============
add_section_heading('Block 3 — Activity 2: Adversarial Sensor Challenge',
                    'Running time 10:35 to 11:00 · 25 minutes')

# ---- 36
add_slide_header(36, 'Activity 2 Intro', '10:35')
screen("Large title 'Sensor Fault Detection' over a purple slide.")
say("Okay, we're back. Activity 2 is different. This is a competition. You'll work in teams. "
    "The first team to identify all three attacks correctly gets bragging rights for the day.")
say("Here's the setup. One of engine seventeen's sensors is reporting bad data. You don't know which "
    "sensor. You don't know what kind of attack. Your job is to find both.")

# ---- 37
add_slide_header(37, 'Scenario', '10:36')
screen("A short story-style message from maintenance ops in a yellow callout box.")
say("You're the data science team at an airline. Maintenance ops just sent you this message.")
do("Read the callout box on the slide aloud.")
say("'Something is wrong with engine seventeen. The pattern looks like a sensor drift, a frozen sensor, "
    "or noise spikes, but we can't tell which channel is at fault. Localize the faulty sensor before "
    "the next flight.'")
say("You have two tools. Your Soft Decision Tree from Activity 1, which is explainable. And we'll add "
    "a RandomForest baseline, which is more accurate but opaque. Use both. See which one is actually useful.")

# ---- 38
add_slide_header(38, 'Teams', '10:37')
screen("Bullet list of team-formation rules and a centered line that says 'First team to identify all three correctly wins.'")
say("Form groups of three to four people. Pick the people next to you. Thirty seconds.")
do("Walk to the back of the room while teams form. Come back.")
say("Use one notebook per team. One person types, the others think and argue. Once we start the clock, "
    "you have seventeen minutes for analysis. First team to identify all three correctly — sensor and "
    "attack type for each file — wins.")

# ---- 39
add_slide_header(39, 'Three Attack Files', '10:38')
screen("A three-row table mapping attack_A, attack_B, attack_C to attack types.")
do("Point at the table row by row.")
say("Three files. Attack A is a drift — a constant offset added to one sensor. Attack B is stuck-at — "
    "one sensor frozen at a single value. Attack C is Gaussian noise — random noise injected into "
    "one sensor.")
say("All three files are copies of clean engine seventeen data. In each one, exactly one sensor channel "
    "has been manipulated. Find which one.")

# ---- 40 to 46 are reference for students; teacher mostly stays on the clock slide
add_slide_header(40, 'A2 Step 1 — RF Baseline (reference)', '10:40')
screen("STEP 1 / 7. Code adds a RandomForest. Right side shows two near-identical accuracies.")
do("Briefly walk through the slide if a team asks. Most teams will just read it in their notebook.")
say("Step 1. Add a RandomForest baseline. SoftTree about 0.841, RandomForest about 0.843. Essentially "
    "identical on clean data. The point of this activity is what happens OFF clean data.")

add_slide_header(41, 'A2 Step 2 — Clean Baseline (reference)', '10:42')
screen("STEP 2 / 7. Code predicts on clean engine 17. Right side: 276 cycles, 0.95 agreement.")
say("Step 2. Predict on clean engine seventeen. Both models agree on ninety-five percent of cycles. "
    "That's our baseline.")

add_slide_header(42, 'A2 Step 3 — Score Attacks (reference)', '10:44')
screen("STEP 3 / 7. Loop counting prediction changes per attack file. Right side: per-file Tree and RF counts.")
say("Step 3. Run both models on each attack file. Count cycle-by-cycle disagreement against the clean "
    "baseline. Watch Attack C — RandomForest changes only one cycle. It is essentially blind to the attack. "
    "The Soft Tree caught eight.")

add_slide_header(43, 'A2 Step 4 — Attack A (reference)', '10:48')
screen("STEP 4 / 7. Plot loop. Right side: 4x4 grid showing clean (blue) and attack (orange) overlays.")
say("Step 4. Visualize Attack A. Fourteen panels overlap. One shows two parallel lines — a constant "
    "offset. That's drift. Answer: sensor eleven.")

add_slide_header(44, 'A2 Step 5 — Attack B (reference)', '10:51')
screen("STEP 5 / 7. Same plot pattern; Attack B's one sensor is a flat orange line.")
say("Step 5. Attack B. One sensor's orange line is completely flat while the blue line keeps evolving. "
    "Stuck-at. Sensor fourteen — which was also a pillar in Activity 1 step 10. Attackers went for "
    "what the model relies on the most.")

add_slide_header(45, 'A2 Step 6 — Attack C (reference)', '10:54')
screen("STEP 6 / 7. Attack C plot; one sensor's orange line is noticeably jaggier than blue.")
say("Step 6. Attack C. Same trend, much higher variance. Gaussian noise. Sensor nine. The trend is "
    "preserved, only the noise level changes — that's why RandomForest kept saying 'fine'.")

add_slide_header(46, 'A2 Step 7 — Quantify (reference)', '10:57')
screen("STEP 7 / 7. Loop printing top-3 divergent sensors per attack. Right side: clean numerical confirmation.")
say("Step 7. Quantify the divergence. For each attack, only one sensor has a non-zero divergence. "
    "Sensor eleven for A, sensor fourteen for B, sensor nine for C. The other fourteen are identical "
    "between clean and attack. Attackers manipulate one sensor and try to stay hidden. Our model "
    "caught them anyway.")


doc.add_page_break()


# ============ BLOCK 4 ============
add_section_heading('Block 4 — Live Answers and Discussion',
                    'Running time 11:00 to 11:15 · 15 minutes')

# ---- 47
add_slide_header(47, 'Time on Clock — Teams Working', '10:40')
screen("Large '⏱ 17 minutes' on a dark slide. This is the slide that stays up while teams work.")
do("This is the slide the projector shows while teams work. The seven step slides above are reference "
   "material the students follow IN THEIR NOTEBOOKS. You do NOT need to walk through each one on stage. "
   "Just leave this clock slide up.")
say("Seventeen minutes on the clock starting now. The seven steps are in your notebook — follow them "
    "top to bottom. Discuss with your team. I'm walking around. Raise a hand if you have questions.")
say("I'll give hints, but I will not give you the answer. Go.")
do("Start the timer. Walk the room. Provoke teams who finish early: 'Show me your reasoning, not just "
   "your answer.' Help teams that are stuck — but don't reveal which sensor.")

# ---- 48
add_slide_header(48, 'Answers', '11:00')
screen("A three-row answer table: A=11 drift, B=14 stuck-at, C=9 noise, plus a trophy emoji.")
say("Time's up. Here are the answers.")
do("Point at each row in the table.")
say("Attack A is sensor eleven. Drift. Attack B is sensor fourteen. Stuck-at. Attack C is sensor nine. "
    "Gaussian noise.")
say("Show of hands — which team got all three?")
do("Acknowledge the winning team by name if you can. Lead a round of applause.")

# ---- 49
add_slide_header(49, 'What Just Happened', '11:03')
screen("Three bullets explaining the teaching point, then a big italic line about explainability as a "
       "debugging tool.")
say("Here's what we just demonstrated.")
do("Point at the first bullet.")
say("The RandomForest changed its prediction under attack — but it couldn't tell you which sensor "
    "caused the change.")
do("Point at the second bullet.")
say("The Soft Decision Tree's split weights shifted in a sensor-specific way. That shift IS the "
    "explanation.")
do("Point at the third bullet.")
say("It's a fingerprint of which sensor the model is leaning on.")
say("The lesson: explainability isn't just for regulatory compliance. It's also a debugging tool. The "
    "same property that lets you defend a model to the FAA also helps your maintenance engineers find "
    "the broken hardware.")


doc.add_page_break()


# ============ BLOCK 5 ============
add_section_heading('Block 5 — GitHub Contribution Sprint',
                    'Running time 11:15 to 11:30 · 15 minutes')

# ---- 50
add_slide_header(50, 'Sprint Intro', '11:15')
screen("Green slide. Big title with wrench emoji: 'GitHub Contribution Sprint'.")
say("Last fifteen minutes. This part is different from everything else. You're going to make a real "
    "open-source contribution to a real Python library that's on PyPI. Not a workshop exercise. A real "
    "pull request that I will review and merge live, in front of you, right now. Your name on it. Forever.")

# ---- 51
add_slide_header(51, 'Sprint Mechanics', '11:17')
screen("A seven-step numbered list of how to fork, edit in the GitHub web editor, and open a PR.")
say("Here's how it works. Seven steps.")
do("Point at each numbered step as you say it.")
say("One. Open github.com/cgrtml/neural-trees. Go to the Issues tab.")
say("Two. Filter by 'good first issue'. Pick one — they're all ten to twenty minute scope.")
say("Three. Comment 'I'm taking this' so two people don't pick the same one.")
say("Four. Click the Fork button — top right of the page. That gives you your own copy of the repo.")
say("Five. Open the file you want to edit — in YOUR fork, not in mine.")
do("Pause briefly. This next part is the trick that saves them ten minutes.")
say("And here's the trick: press the dot key on your keyboard. Just the period key. GitHub's full web "
    "editor opens, right there in the browser. You can edit like you're in VS Code.")
say("Six. Commit your change on a new branch. Click 'Compare and pull request'. Add a short description. "
    "Submit.")
say("Seven. I'll review it from my laptop right here on stage. If it's correct, I merge it on the spot.")
say("No local git. No clone. No tokens. Everything in the browser. About twenty issues are waiting. "
    "Pick what fits your level.")

# ---- 52
add_slide_header(52, 'Why Do It', '11:20')
screen("Three bullets explaining the value: profile credit, real package, named in README.")
say("Why is this worth fifteen minutes? Three reasons.")
do("Point at bullet one.")
say("Your contribution shows up on your GitHub profile. That's visible to recruiters, internship "
    "coordinators, grad school admissions. Much harder to fake than a tutorial certificate.")
do("Point at bullet two.")
say("You become part of a Python package that other people install with pip. Your code runs on their "
    "machines.")
do("Point at bullet three.")
say("Every merged PR will be credited in the neural-trees README under a section called WSU Workshop "
    "Contributors. That section will exist after today, with your names on it.")
say("Open the repo. Pick an issue. Go.")
do("Walk to your laptop on stage. Open the PRs tab. Start reviewing as PRs come in. Merge live where possible.")


doc.add_page_break()


# ============ BLOCK 6 ============
add_section_heading('Block 6 — Wrap-up and Q&A',
                    'Running time 11:30 to 11:45 · 15 minutes')

# ---- 53
add_slide_header(53, 'EU AI Act', '11:30')
screen("Three articles of the EU AI Act listed, plus a footer mentioning SR 11-7, FDA, FAA.")
say("Quick context for why all of this matters in the real world. The EU AI Act took effect in 2025. "
    "If your AI system serves any European customer, three articles apply.")
do("Point at each article.")
say("Article 13: transparency. Article 14: human oversight. Article 15: robustness against adversarial "
    "inputs and faults. The exact things we did today.")
say("And the US is moving in the same direction. SR 11-7 for banking. FDA for medical devices. FAA for "
    "aviation. The era of 'just ship the LSTM' is ending. The era of 'show your work' is starting.")

# ---- 54
add_slide_header(54, 'Wrap-Up', '11:33')
screen("Three numbered takeaways on a green slide.")
say("Three things to take with you.")
do("Point at point one.")
say("In safety-critical AI, explainability is not optional. It's a regulatory requirement.")
do("Point at point two.")
say("Soft Decision Trees give you both worlds. Neural-network-level accuracy with tree-level interpretability.")
do("Point at point three.")
say("You actually did this. Not me. Not a video. You. You trained a model, you localized a sensor fault, "
    "and many of you just shipped your first open-source contribution.")

# ---- 55
add_slide_header(55, 'Career Bridge', '11:36')
screen("A short paragraph about Hezarfen LLC and a centered 'reach out.' line.")
say("Where does this work go next? Banks, insurers, healthcare, aerospace. Every regulated industry "
    "needs explainable AI right now. My company Hezarfen builds compliance-ready AI for US mid-market "
    "banks under SR 11-7 and the EU AI Act.")
say("If anything you saw today interests you — a thesis topic, an open-source contribution, internship "
    "questions, or just a conversation about where to go after graduation — reach out. My email and "
    "LinkedIn are on the next slide.")

# ---- 56
add_slide_header(56, 'Links', '11:39')
screen("Centered block with name, three links: GitHub, LinkedIn, workshop repo.")
say("Three things worth bookmarking. My GitHub — where neural-trees lives. My LinkedIn — happy to "
    "connect with any of you. And the workshop repo on GitHub. Everything we did today is there: the "
    "data, the notebooks, the slides.")
do("Pause. Let students take a photo of the slide.")
say("Take a photo if you want.")

# ---- 57
add_slide_header(57, 'Thank You', '11:42')
screen("Large 'Thank you' on a dark slide. Smaller text thanking Dr. Sergey Lapin and Jeremy.")
say("Thank you. Special thanks to Dr. Sergey Lapin and to Jeremy for the invitation. Without them this "
    "wouldn't have happened.")
say("I'll stay for five minutes after the official close. If you have questions, come up to the stage. "
    "If you want to follow up later, you have my email. Have a great rest of the week.")


doc.save(OUT_PATH)
print(f"Saved: {OUT_PATH}")
